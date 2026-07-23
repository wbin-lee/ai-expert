#!/usr/bin/env python3
"""Generate Word submission with project topic and introduction only."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).parent / "현업과제_제안서_ClaudeCoach.docx"

INTRO_PROPOSAL = (
    "최근 사내 Claude Code Enterprise가 도입되었으며, 1인당 월 약 $180(종량제) "
    "사용 한도가 있습니다. "
    "사용 로그를 분석해 우수·저토큰 고효율 활용 사례를 검색·추천하고, "
    "맞춤 코칭을 제공하는 ClaudeCoach 시스템을 구축합니다. "
    "기존 활용 대시보드와 AWS 게이트웨이를 확장하며 코칭 기능(MCP/게이트웨이 플러그인)을 "
    "추가합니다. "
    "① 로그 익명 정제·토큰/비용 메타·우수 세션 라벨링, "
    "② PatternEmbed(Contrastive FT), "
    "③ CoachModel(LoRA SFT), "
    "④ BudgetCoach(한도 대비 사용률·저토큰 사례·컨텍스트 축소 권고), "
    "⑤ DS LLM API·게이트웨이 연동(suggest_usage, find_pattern)으로 구성됩니다. "
    "개발자가 「예산 70% 썼는데 레거시 분석 어떻게 해?」라고 질문하면 "
    "저토큰 우수 세션·작업 분할·MCP 활용 패턴을 제안합니다."
)

INTRO_TECH = (
    "지도교수: 서울대 윤성로 교수님. "
    "7월 베이스 모델 선정 PoC·골드셋 확보 후 8월 본격 수행. "
    "학습(임베딩 FT, LoRA SFT)은 AI 포털 GPU, 운영은 DS LLM API·검색 컨텍스트. "
    "교수님 지도: 활용·비용 효율 패턴 분석, 우수 세션 정의, 실험·평가 설계. "
    "평가는 Hit@k, 코칭 유용도 설문, DS 베이스 LLM 대비 A/B, (탐색) 동일 태스크 토큰 절감."
)

INTRO_EFFECT = (
    "Claude 사용자의 활용도·정착·생산성과 월 $180 한도 내 Bedrock 비용 효율 향상에 기여합니다. "
    "10월 MVP: 코칭 E2E 데모, 파일럿 50회+ 호출, Hit@3 60%+, 한도 사용률 대시보드 연동. "
    "AI-Expert 요건(모델 선정 → Fine-tuning → 서빙 연동)을 충족합니다."
)


def set_run_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=(30, 60, 110))


def add_labeled_paragraph(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    label_run = p.add_run(label)
    set_run_font(label_run, bold=True)
    body_run = p.add_run(body)
    set_run_font(body_run)


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(11)

    add_section_heading(doc, "1. 프로젝트 주제")
    topic = doc.add_paragraph()
    topic.paragraph_format.space_after = Pt(12)
    topic_run = topic.add_run(
        "사내 AI 코딩 도구 활용·예산 효율 코칭 시스템 (ClaudeCoach)"
    )
    set_run_font(topic_run, size=11, bold=True)

    add_section_heading(doc, "2. 프로젝트 소개")
    add_labeled_paragraph(doc, "제안 내용  ", INTRO_PROPOSAL)
    add_labeled_paragraph(doc, "기술·인프라  ", INTRO_TECH)
    add_labeled_paragraph(doc, "기대 효과  ", INTRO_EFFECT)

    total = len(INTRO_PROPOSAL + INTRO_TECH + INTRO_EFFECT)
    doc.save(OUTPUT)
    print(f"소개 본문 글자수: {total}")
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()